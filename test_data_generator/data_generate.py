import os
import csv
import xml.etree.ElementTree as ET
from faker import Faker
from docx import Document
from PIL import Image, ImageDraw
from fpdf import FPDF

def create_test_data(file_count, file_format, text_length=0):
    if not os.path.exists("test_data"):
        os.makedirs("test_data")

    fake = Faker()

    for i in range(file_count):
        filename = f"test_data/test_file_{i + 1}.{file_format}"
        
        if file_format == "txt":
            text = fake.paragraphs(nb=text_length)
            with open(filename, "w") as file:
                file.write("\n".join(text))
        elif file_format == "docx":
            document = Document()
            text = fake.paragraphs(nb=text_length)
            for paragraph in text:
                document.add_paragraph(paragraph)
            document.save(filename)
        elif file_format == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            text = fake.paragraphs(nb=text_length)
            pdf.multi_cell(0, 10, "\n".join(text))
            pdf.output(filename)
        elif file_format == "csv":
            with open(filename, "w", newline="") as file:
                csv_writer = csv.writer(file)
                for _ in range(text_length):
                    csv_writer.writerow([fake.name(), fake.address(), fake.random_int()])
        elif file_format == "xml":
            root = ET.Element("data")
            for _ in range(text_length):
                item = ET.SubElement(root, "item")
                ET.SubElement(item, "name").text = fake.name()
                ET.SubElement(item, "address").text = fake.address()
                ET.SubElement(item, "number").text = str(fake.random_int())
            tree = ET.ElementTree(root)
            tree.write(filename)
        elif file_format == "jpg":
            img = Image.new("RGB", (800, 600), color="white")
            d = ImageDraw.Draw(img)
            d.text((10, 10), "Test Image no: "+str(i+1), fill="black")
            img.save(filename)

if __name__ == "__main__":
    file_count = int(input("Enter the number of files to generate: "))
    file_format = input("Enter the file format (txt, docx, pdf, csv, xml, jpg): ").lower()
    
    if file_format not in ["txt", "docx", "pdf", "csv", "xml", "jpg"]:
        print("Invalid file format specified. Please provide one of these \n ['txt', 'docx', 'csv', 'xml', 'jpg', 'pdf']")
    else:
        if file_format != "jpg":
            text_length = int(input("Enter the number of items/paragraphs to generate in each file: "))
            create_test_data(file_count, file_format, text_length)
        else :
            create_test_data(file_count, file_format)
        print(f"{file_count} {file_format} files generated in the 'test_data' directory.")